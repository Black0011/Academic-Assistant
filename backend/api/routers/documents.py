"""Knowledge document API — upload arbitrary docs for RAG (M7.3).

Surface (PLAN §20.8 M7.3):

* ``POST   /api/documents/ingest``                multipart **or** JSON
* ``GET    /api/documents``
* ``GET    /api/documents/{doc_id}``
* ``GET    /api/documents/{doc_id}/chunks``
* ``POST   /api/documents/{doc_id}:reindex``      re-chunk + rebuild vectors
* ``DELETE /api/documents/{doc_id}``              cascades to vector store
* ``POST   /api/documents/search``                vector RAG over docs only

Ingest accepts:

* ``multipart/form-data``: a ``file`` field (pdf / md / txt) plus optional
  ``title``, ``tags``, ``source_kind``, ``source_uri`` text fields.
* ``application/json``: ``{title, raw_text, source_kind, ...}``. Useful
  for clipboard / programmatic ingest.

The pipeline is deterministic: chunk_id = ``f"{doc_id}#{idx:04d}"`` so
re-indexing the same document never breaks links from prior runs.
"""

from __future__ import annotations

from datetime import UTC, datetime
from typing import Any, Literal

import structlog
from fastapi import (
    APIRouter,
    Depends,
    HTTPException,
    Query,
    Request,
    Response,
)
from pydantic import BaseModel, ConfigDict, Field
from pydantic import ValidationError as PydanticValidationError

from backend.core.app_state import AppState, get_app_state
from backend.core.text import pdf_to_markdown
from backend.memory.base import DocumentStore, stable_id
from backend.memory.chunker import chunk_markdown
from backend.memory.document_store import (
    heuristic_summary,
    make_chunk_id,
)
from backend.memory.models import (
    DocChunk,
    DocChunkHit,
    DocumentSourceKind,
    KnowledgeDocument,
)

log = structlog.get_logger(__name__)


MAX_INGEST_BYTES = 25 * 1024 * 1024  # 25 MB, matches manuscripts upload cap.

router = APIRouter(prefix="/api/documents", tags=["documents"])


# ---------------------------------------------------------------------------
# Request / response DTOs
# ---------------------------------------------------------------------------


class IngestDocumentJSONInput(BaseModel):
    """Payload for ``application/json`` ingest."""

    model_config = ConfigDict(extra="forbid")

    title: str = Field("", description="Auto-derived from raw_text/heading when empty.")
    raw_text: str = Field(..., min_length=1)
    source_kind: DocumentSourceKind = "note"
    source_uri: str = ""
    summary: str = ""
    tags: list[str] = Field(default_factory=list)
    user_id: str | None = None
    session_id: str | None = None
    source_run_id: str | None = None
    target_tokens: int = Field(800, ge=100, le=4000)
    overlap_tokens: int = Field(100, ge=0, le=1000)


class IngestResultDTO(BaseModel):
    document: KnowledgeDocument
    chunks_indexed: int
    indexer_ms: int = 0


class DocumentListResponse(BaseModel):
    items: list[KnowledgeDocument]
    total: int


class DocumentChunkPage(BaseModel):
    items: list[DocChunk]
    total: int


class SearchInput(BaseModel):
    model_config = ConfigDict(extra="forbid")

    q: str = Field(..., min_length=1)
    top_k: int = Field(5, ge=1, le=50)
    filters: dict[str, Any] = Field(default_factory=dict)


class SearchResponse(BaseModel):
    items: list[DocChunkHit]
    total: int


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _require_documents(state: AppState) -> DocumentStore:
    if state.memory is None or state.memory.documents is None:
        raise HTTPException(status_code=503, detail="document store not ready")
    return state.memory.documents


def _split_csv(value: str | None) -> list[str]:
    if not value:
        return []
    return [item.strip() for item in value.split(",") if item.strip()]


def _derive_doc_id(*, title: str, source_uri: str, raw_text: str) -> str:
    seed = (title or "").strip().lower() or raw_text[:64]
    return stable_id("doc", source_uri or "", seed)


def _derive_title(raw_text: str, *, fallback: str) -> str:
    """Pick a title from the first heading / first non-blank line."""
    for line in raw_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("#"):
            return stripped.lstrip("# ").strip() or fallback
        return stripped[:120]
    return fallback


def _decode_blob(raw: bytes, *, content_type: str, filename: str) -> tuple[str, dict[str, Any]]:
    """Best-effort decode of an uploaded file into markdown."""
    name = (filename or "").lower()
    ct = (content_type or "").lower()
    if "pdf" in ct or name.endswith(".pdf") or raw.startswith(b"%PDF"):
        body, meta = pdf_to_markdown(raw)
        return body, {"format": "pdf", **meta}
    if name.endswith(".docx") or "vnd.openxmlformats-officedocument.wordprocessingml" in ct:
        body = _docx_to_text(raw)
        return body, {"format": "docx"}
    try:
        return raw.decode("utf-8"), {"format": "text"}
    except UnicodeDecodeError:
        return raw.decode("utf-8", errors="replace"), {"format": "text", "decode": "replace"}


def _docx_to_text(raw: bytes) -> str:
    """Extract text from a .docx file (Word 2007+)."""
    from io import BytesIO
    try:
        from docx import Document
        doc = Document(BytesIO(raw))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    except Exception:
        return raw.decode("utf-8", errors="replace")


def _classify_source(content_type: str, filename: str) -> DocumentSourceKind:
    name = (filename or "").lower()
    ct = (content_type or "").lower()
    if "pdf" in ct or name.endswith(".pdf"):
        return "pdf_upload"
    if name.endswith(".md") or name.endswith(".markdown") or "markdown" in ct:
        return "md_upload"
    if name.endswith(".docx") or "wordprocessingml" in ct:
        return "docx_upload"
    return "txt_upload"


# ---------------------------------------------------------------------------
# Ingest
# ---------------------------------------------------------------------------


def _now_ms() -> int:
    return int(datetime.now(UTC).timestamp() * 1000)


async def _ingest_pipeline(
    *,
    store: DocumentStore,
    title: str,
    raw_text: str,
    source_kind: DocumentSourceKind,
    source_uri: str,
    summary: str,
    tags: list[str],
    user_id: str | None,
    session_id: str | None,
    source_run_id: str | None,
    extras: dict[str, Any],
    target_tokens: int,
    overlap_tokens: int,
) -> IngestResultDTO:
    if not raw_text.strip():
        raise HTTPException(status_code=400, detail="raw_text is empty")

    final_title = (title or "").strip() or _derive_title(raw_text, fallback="Untitled document")
    doc_id = _derive_doc_id(title=final_title, source_uri=source_uri, raw_text=raw_text)
    raw_chunks = chunk_markdown(
        raw_text,
        target_tokens=target_tokens,
        overlap_tokens=overlap_tokens,
    )
    if not raw_chunks:
        raise HTTPException(status_code=400, detail="document produced zero chunks")

    chunks = [
        DocChunk(
            chunk_id=make_chunk_id(doc_id, idx),
            doc_id=doc_id,
            idx=idx,
            text=raw.text,
            char_offset_start=raw.char_offset_start,
            char_offset_end=raw.char_offset_end,
            section_path=list(raw.section_path),
            tags=[],
        )
        for idx, raw in enumerate(raw_chunks)
    ]

    document = KnowledgeDocument(
        doc_id=doc_id,
        title=final_title,
        source_kind=source_kind,
        source_uri=(source_uri.strip() or None) if source_uri else None,
        summary=(summary.strip() or heuristic_summary(raw_text)),
        raw_text=raw_text,
        tags=[t for t in {*tags} if t],
        chunk_ids=[c.chunk_id for c in chunks],
        bytes=len(raw_text.encode("utf-8")),
        user_id=user_id,
        session_id=session_id,
        source_run_id=source_run_id,
        extras=extras,
    )

    started = _now_ms()
    await store.write(document, chunks)
    elapsed = _now_ms() - started

    final = await store.get(doc_id)
    if final is None:  # pragma: no cover - defensive
        raise HTTPException(status_code=500, detail="write did not persist")
    return IngestResultDTO(
        document=final,
        chunks_indexed=len(chunks),
        indexer_ms=elapsed,
    )


@router.post(
    "/ingest",
    response_model=IngestResultDTO,
    status_code=201,
    summary="Ingest a document (PDF / MD / TXT or pure JSON) and index its chunks.",
)
async def ingest_document(
    request: Request,
    state: AppState = Depends(get_app_state),
) -> IngestResultDTO:
    store = _require_documents(state)
    content_type = (request.headers.get("content-type") or "").lower()

    if "multipart/form-data" in content_type:
        form = await request.form()
        file = form.get("file")
        if file is None or not hasattr(file, "read"):
            raise HTTPException(status_code=400, detail="multipart ingest requires a 'file' field")
        raw = await file.read()
        if not isinstance(raw, bytes) or len(raw) == 0:
            raise HTTPException(status_code=400, detail="empty upload")
        if len(raw) > MAX_INGEST_BYTES:
            raise HTTPException(status_code=413, detail="file too large")

        def _str(name: str) -> str:
            v = form.get(name)
            return v if isinstance(v, str) else ""

        body_text, extras = _decode_blob(
            raw,
            content_type=getattr(file, "content_type", "") or "",
            filename=getattr(file, "filename", "") or "",
        )
        source_kind = _classify_source(
            getattr(file, "content_type", "") or "",
            getattr(file, "filename", "") or "",
        )
        title = _str("title") or (getattr(file, "filename", "") or "")
        return await _ingest_pipeline(
            store=store,
            title=title,
            raw_text=body_text,
            source_kind=source_kind,
            source_uri=_str("source_uri"),
            summary=_str("summary"),
            tags=_split_csv(_str("tags")),
            user_id=_str("user_id") or None,
            session_id=_str("session_id") or None,
            source_run_id=_str("source_run_id") or None,
            extras=extras,
            target_tokens=int(_str("target_tokens") or 800),
            overlap_tokens=int(_str("overlap_tokens") or 100),
        )

    if "application/json" in content_type:
        try:
            payload = await request.json()
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"invalid json body: {exc}") from exc
        try:
            parsed = IngestDocumentJSONInput.model_validate(payload)
        except PydanticValidationError as exc:
            raise HTTPException(status_code=422, detail=exc.errors()) from exc
        return await _ingest_pipeline(
            store=store,
            title=parsed.title,
            raw_text=parsed.raw_text,
            source_kind=parsed.source_kind,
            source_uri=parsed.source_uri,
            summary=parsed.summary,
            tags=list(parsed.tags),
            user_id=parsed.user_id,
            session_id=parsed.session_id,
            source_run_id=parsed.source_run_id,
            extras={"format": "manual"},
            target_tokens=parsed.target_tokens,
            overlap_tokens=parsed.overlap_tokens,
        )

    raise HTTPException(
        status_code=415,
        detail=(
            "Content-Type must be 'multipart/form-data' (with a 'file' field) "
            "or 'application/json' (with raw_text)"
        ),
    )


# ---------------------------------------------------------------------------
# Read endpoints
# ---------------------------------------------------------------------------


@router.get(
    "",
    response_model=DocumentListResponse,
    summary="List ingested documents (most recent first).",
)
async def list_documents(
    user_id: str | None = Query(None),
    tag: str | None = Query(None),
    limit: int = Query(50, ge=1, le=500),
    offset: int = Query(0, ge=0),
    state: AppState = Depends(get_app_state),
) -> DocumentListResponse:
    store = _require_documents(state)
    items = await store.list_all(user_id=user_id)
    if tag:
        items = [doc for doc in items if tag in doc.tags]
    page = items[offset : offset + limit]
    return DocumentListResponse(items=page, total=len(items))


@router.get(
    "/{doc_id}",
    response_model=KnowledgeDocument,
    summary="Get one document (full raw text included).",
)
async def get_document(
    doc_id: str,
    state: AppState = Depends(get_app_state),
) -> KnowledgeDocument:
    store = _require_documents(state)
    doc = await store.get(doc_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="document not found")
    return doc


@router.get(
    "/{doc_id}/chunks",
    response_model=DocumentChunkPage,
    summary="List chunks for one document (paginated).",
)
async def list_chunks(
    doc_id: str,
    offset: int = Query(0, ge=0),
    limit: int = Query(100, ge=1, le=1000),
    state: AppState = Depends(get_app_state),
) -> DocumentChunkPage:
    store = _require_documents(state)
    doc = await store.get(doc_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="document not found")
    chunks = await store.get_chunks(doc_id)
    return DocumentChunkPage(items=chunks[offset : offset + limit], total=len(chunks))


# ---------------------------------------------------------------------------
# Mutations
# ---------------------------------------------------------------------------


class ReindexInput(BaseModel):
    model_config = ConfigDict(extra="forbid")

    target_tokens: int = Field(800, ge=100, le=4000)
    overlap_tokens: int = Field(100, ge=0, le=1000)


@router.post(
    "/{doc_id}:reindex",
    response_model=IngestResultDTO,
    summary="Re-chunk an existing document and rebuild its vector entries.",
)
async def reindex_document(
    doc_id: str,
    body: ReindexInput | None = None,
    state: AppState = Depends(get_app_state),
) -> IngestResultDTO:
    store = _require_documents(state)
    existing = await store.get(doc_id)
    if existing is None:
        raise HTTPException(status_code=404, detail="document not found")
    params = body or ReindexInput(target_tokens=800, overlap_tokens=100)
    return await _ingest_pipeline(
        store=store,
        title=existing.title,
        raw_text=existing.raw_text,
        source_kind=existing.source_kind,
        source_uri=existing.source_uri or "",
        summary=existing.summary,
        tags=list(existing.tags),
        user_id=existing.user_id,
        session_id=existing.session_id,
        source_run_id=existing.source_run_id,
        extras={"format": "reindex", **dict(existing.extras)},
        target_tokens=params.target_tokens,
        overlap_tokens=params.overlap_tokens,
    )


@router.delete(
    "/{doc_id}",
    status_code=204,
    summary="Delete a document and prune its chunks from the vector store.",
)
async def delete_document(
    doc_id: str,
    state: AppState = Depends(get_app_state),
) -> Response:
    store = _require_documents(state)
    deleted = await store.delete(doc_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="document not found")
    return Response(status_code=204)


# ---------------------------------------------------------------------------
# P14.B — PATCH metadata
#
# Why a separate endpoint instead of overloading ingest / reindex?
# ----------------------------------------------------------------
# - ingest produces a NEW doc_id (we can't override that path mid-edit).
# - reindex re-chunks + re-embeds; for "fix the title typo" that's
#   ~100ms→several-seconds for a 100-page PDF, which is wasteful and
#   surprising.
# A dedicated PATCH that NEVER touches raw_text (and therefore doesn't
# re-chunk / re-embed) keeps the edit path cheap and predictable.
#
# raw_text is intentionally absent from the editable surface — see the
# extended docstring on ``DocumentStore.update_metadata``.
# ---------------------------------------------------------------------------


class UpdateDocumentInput(BaseModel):
    model_config = ConfigDict(extra="forbid")

    title: str | None = Field(None, min_length=1, max_length=400)
    summary: str | None = Field(None, max_length=4000)
    tags: list[str] | None = None
    source_kind: DocumentSourceKind | None = None
    source_uri: str | None = Field(None, max_length=2000)


@router.patch(
    "/{doc_id}",
    response_model=KnowledgeDocument,
    summary="Edit document metadata (title / summary / tags / source). Does NOT re-chunk or re-embed.",
)
async def update_document_metadata(
    doc_id: str,
    body: UpdateDocumentInput,
    state: AppState = Depends(get_app_state),
) -> KnowledgeDocument:
    store = _require_documents(state)
    updated = await store.update_metadata(
        doc_id,
        title=body.title,
        summary=body.summary,
        tags=body.tags,
        source_kind=body.source_kind,
        source_uri=body.source_uri,
    )
    if updated is None:
        raise HTTPException(status_code=404, detail="document not found")
    return updated


# ---------------------------------------------------------------------------
# Search
# ---------------------------------------------------------------------------


@router.post(
    "/search",
    response_model=SearchResponse,
    summary="Vector RAG search restricted to ingested documents.",
)
async def search_documents(
    body: SearchInput,
    state: AppState = Depends(get_app_state),
) -> SearchResponse:
    store = _require_documents(state)
    where = dict(body.filters or {})
    hits = await store.search_chunks(body.q, k=body.top_k, where=where)
    return SearchResponse(items=hits, total=len(hits))


# Re-export the type used by other routers (kept here so tests import it
# from a single canonical location).
SourceKind = Literal["pdf_upload", "md_upload", "txt_upload", "note", "url", "clipboard"]


__all__ = [
    "MAX_INGEST_BYTES",
    "SourceKind",
    "router",
]
